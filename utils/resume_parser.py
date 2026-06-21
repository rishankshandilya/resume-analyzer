"""
resume_parser.py
-----------------
Extracts raw text from uploaded resume files (PDF or DOCX).

Why this is its own module:
The rest of the app (ATS scoring, skill extraction) only cares about
plain text. It shouldn't need to know HOW that text was obtained.
This is the Single Responsibility Principle - one module, one job.
"""

import pdfplumber
import docx
import io


class ParsingError(Exception):
    """Raised when a file can't be read or contains no extractable text."""
    pass


class ResumeParser:
    """
    Object-oriented wrapper around file parsing.
    Usage:
        parser = ResumeParser(uploaded_file)
        text = parser.extract_text()
    """

    SUPPORTED_TYPES = {"pdf", "docx"}

    def __init__(self, uploaded_file):
        """
        uploaded_file: a Streamlit UploadedFile object (has .name and behaves like a file).
        """
        self.uploaded_file = uploaded_file
        self.file_name = uploaded_file.name
        self.file_type = self._detect_type()

    def _detect_type(self) -> str:
        ext = self.file_name.lower().split(".")[-1]
        if ext not in self.SUPPORTED_TYPES:
            raise ParsingError(
                f"Unsupported file type '.{ext}'. Please upload a PDF or DOCX file."
            )
        return ext

    def extract_text(self) -> str:
        """Routes to the correct extraction method based on file type."""
        if self.file_type == "pdf":
            text = self._extract_from_pdf()
        else:
            text = self._extract_from_docx()

        cleaned = self._clean_text(text)

        if not cleaned or len(cleaned) < 30:
            raise ParsingError(
                "Could not extract readable text from this file. "
                "It may be a scanned image rather than a text-based document."
            )
        return cleaned

    def _extract_from_pdf(self) -> str:
        """Uses pdfplumber to pull text from every page of a PDF."""
        text_parts = []
        try:
            # uploaded_file is a file-like object; pdfplumber can read it directly
            with pdfplumber.open(self.uploaded_file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
        except Exception as e:
            raise ParsingError(f"Failed to read PDF file: {e}")

        return "\n".join(text_parts)

    def _extract_from_docx(self) -> str:
        """Uses python-docx to pull text from every paragraph (and table cell) of a DOCX."""
        text_parts = []
        try:
            file_bytes = io.BytesIO(self.uploaded_file.read())
            document = docx.Document(file_bytes)

            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Many resumes use tables for layout (e.g., skills in columns)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

        except Exception as e:
            raise ParsingError(f"Failed to read DOCX file: {e}")

        return "\n".join(text_parts)

    @staticmethod
    def _clean_text(text: str) -> str:
        """Normalizes whitespace so downstream regex/keyword matching is reliable."""
        if not text:
            return ""
        lines = [line.strip() for line in text.splitlines()]
        lines = [line for line in lines if line]  # drop empty lines
        return "\n".join(lines)
